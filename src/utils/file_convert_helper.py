import os
from PIL import Image
from docx import Document
import pandas as pd
import fitz
import requests


class FileConvertHelper:
    def __init__(self, file_url):
        self.file_url = file_url

    def convert_image(self, input_file, output_file, output_format=None):
        if output_format == "jpg" or output_format == "jpeg":
            output_format = "JPEG"
        with Image.open(input_file) as img:
            if img.mode == 'RGBA':
                rgb_img = Image.new("RGB", img.size, (255, 255, 255))
                rgb_img.paste(img, mask=img.split()[3])
                img = rgb_img
            if output_format:
                img.save(output_file, quality=95, format=output_format)
            else:
                img.save(output_file, quality=95)

    def pdf_to_docx(self, pdf_file, docx_file):
        document = Document()
        with fitz.Document(pdf_file) as pdf:
            for page in pdf:
                text = page.get_textpage().extractText()
                document.add_paragraph(text)
        document.save(docx_file)

    def docx_to_markdown(self, docx_file, md_file):
        document = Document(docx_file)
        with open(md_file, "w", encoding="utf-8") as md:
            for para in document.paragraphs:
                md.write(para.text + "\n\n")

    def pdf_to_markdown(self, pdf_file, md_file):
        with fitz.Document(pdf_file) as pdf:
            with open(md_file, "w", encoding="utf-8") as md:
                for page in pdf:
                    text = page.get_textpage().extractText()
                    md.write(text + "\n\n")

    def xlsx_to_csv(self, xlsx_file, csv_file):
        df = pd.read_excel(xlsx_file)
        df.to_csv(csv_file, index=False)

    def csv_to_xlsx(self, csv_file, xlsx_file):
        df = pd.read_csv(csv_file)
        df.to_excel(xlsx_file, index=False)

    def download_file(self, url, folder_path):
        # 确保文件夹存在
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)

        # 从 URL 提取文件名
        file_name = url.split("/")[-1]
        file_path = os.path.join(folder_path, file_name)

        # 发起请求下载文件
        response = requests.get(url)
        if response.status_code == 200:
            with open(file_path, "wb") as file:
                file.write(response.content)
            return file_path
        else:
            raise Exception(f"Error downloading file from {url}")
